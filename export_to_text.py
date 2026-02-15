#!/usr/bin/env python3
"""
Export Project Outline Database to Plain Text File and/or Word Document
"""

import sqlite3
import os
import sys

# ANSI Color codes
class Colors:
    RESET = '\033[0m'
    BOLD = '\033[1m'
    DIM = '\033[2m'
    
    # Foreground colors
    BLACK = '\033[30m'
    RED = '\033[31m'
    GREEN = '\033[32m'
    YELLOW = '\033[33m'
    BLUE = '\033[34m'
    MAGENTA = '\033[35m'
    CYAN = '\033[36m'
    WHITE = '\033[37m'
    
    # Bright foreground colors
    BRIGHT_BLACK = '\033[90m'
    BRIGHT_RED = '\033[91m'
    BRIGHT_GREEN = '\033[92m'
    BRIGHT_YELLOW = '\033[93m'
    BRIGHT_BLUE = '\033[94m'
    BRIGHT_MAGENTA = '\033[95m'
    BRIGHT_CYAN = '\033[96m'
    BRIGHT_WHITE = '\033[97m'
    
    # Background colors
    BG_BLUE = '\033[44m'


def ensure_export_folder():
    """Create document-exports folder if it doesn't exist"""
    export_dir = "document-exports"
    if not os.path.exists(export_dir):
        os.makedirs(export_dir)
        print(f"{Colors.DIM}Created folder: {export_dir}/{Colors.RESET}")
    return export_dir


def list_projects(db_path):
    """List all projects in the database"""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    
    cursor.execute("SELECT id, name, created_at FROM projects ORDER BY created_at DESC")
    projects = cursor.fetchall()
    
    conn.close()
    return projects


def get_project_content(db_path, project_id):
    """Get project content structured for export"""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    
    cursor.execute("SELECT name FROM projects WHERE id = ?", (project_id,))
    project = cursor.fetchone()
    
    if not project:
        conn.close()
        return None, None
    
    project_name = project[0]
    
    # Get structured content
    cursor.execute("""
        SELECT 
            mc.id, mc.name, mc.sort_order,
            sc.id, sc.name, sc.sort_order,
            s.id, s.content, s.sort_order
        FROM major_categories mc
        LEFT JOIN subcategories sc ON mc.id = sc.major_category_id
        LEFT JOIN sentences s ON sc.id = s.subcategory_id
        WHERE mc.project_id = ?
        ORDER BY mc.sort_order, sc.sort_order, s.sort_order
    """, (project_id,))
    
    results = cursor.fetchall()
    conn.close()
    
    # Organize content
    content = {}
    for mc_id, mc_name, mc_order, sc_id, sc_name, sc_order, s_id, s_content, s_order in results:
        if mc_id not in content:
            content[mc_id] = {
                'name': mc_name,
                'order': mc_order,
                'subcategories': {}
            }
        
        if sc_id and sc_id not in content[mc_id]['subcategories']:
            content[mc_id]['subcategories'][sc_id] = {
                'name': sc_name or '',
                'order': sc_order,
                'sentences': []
            }
        
        if s_id and s_content:
            content[mc_id]['subcategories'][sc_id]['sentences'].append(s_content)
    
    return project_name, content


def export_to_text(project_name, content, output_file):
    """Export project to plain text file"""
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(f"PROJECT: {project_name}\n")
        f.write("=" * 80 + "\n\n")
        
        for mc_id in sorted(content.keys(), key=lambda x: content[x]['order']):
            mc_data = content[mc_id]
            f.write(f"{mc_data['name']}\n")
            f.write("-" * len(mc_data['name']) + "\n")
            
            for sc_id in sorted(mc_data['subcategories'].keys(), key=lambda x: mc_data['subcategories'][x]['order']):
                sc_data = mc_data['subcategories'][sc_id]
                
                if sc_data['name']:
                    f.write(f"\n{sc_data['name']}\n")
                
                for sentence in sc_data['sentences']:
                    f.write(f"{sentence}\n\n")
            
            f.write("\n")
    
    return True


def export_to_docx(project_name, content, output_file):
    """Export project to Word document"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        print(f"\n{Colors.YELLOW}Warning:{Colors.RESET} python-docx not installed. Installing now...")
        os.system("sudo pip3 install python-docx")
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    
    doc = Document()
    
    # Set default font and spacing for the document (APA 7 format)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set double spacing for Normal style
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 2.0
    paragraph_format.space_after = Pt(0)
    
    # Add title (APA 7 format)
    title = doc.add_heading(project_name, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.line_spacing = 2.0
    title.paragraph_format.space_after = Pt(0)
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
    
    # Add content
    for mc_id in sorted(content.keys(), key=lambda x: content[x]['order']):
        mc_data = content[mc_id]
        
        # Add major category heading (APA 7 Level 1)
        heading1 = doc.add_heading(mc_data['name'], level=1)
        heading1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heading1.paragraph_format.line_spacing = 2.0
        heading1.paragraph_format.space_after = Pt(0)
        for run in heading1.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
        
        for sc_id in sorted(mc_data['subcategories'].keys(), key=lambda x: mc_data['subcategories'][x]['order']):
            sc_data = mc_data['subcategories'][sc_id]
            
            # Add subcategory heading if it has a name (APA 7 Level 2)
            if sc_data['name']:
                heading2 = doc.add_heading(sc_data['name'], level=2)
                heading2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                heading2.paragraph_format.line_spacing = 2.0
                heading2.paragraph_format.space_after = Pt(0)
                for run in heading2.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.bold = True
            
            # Add sentences (APA 7 double-spaced)
            for sentence in sc_data['sentences']:
                p = doc.add_paragraph(sentence)
                p.paragraph_format.line_spacing = 2.0
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.first_line_indent = Pt(36)  # 0.5 inch indent
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    
    doc.save(output_file)
    return True


def clear_screen():
    """Clear the terminal screen"""
    os.system('clear' if os.name != 'nt' else 'cls')


def get_terminal_size():
    """Get terminal size"""
    try:
        rows, cols = os.popen('stty size', 'r').read().split()
        return int(rows), int(cols)
    except:
        return 24, 80


def print_header(title):
    """Print header bar"""
    rows, cols = get_terminal_size()
    print(f"\n{Colors.BG_BLUE}{Colors.BRIGHT_WHITE}{Colors.BOLD}" + " "*cols + f"{Colors.RESET}")
    header_text = f"  {title}"
    padding = " " * (cols - len(header_text))
    print(f"{Colors.BG_BLUE}{Colors.BRIGHT_WHITE}{Colors.BOLD}{header_text}{padding}{Colors.RESET}")
    print(f"{Colors.BG_BLUE}{Colors.BRIGHT_WHITE}{Colors.BOLD}" + " "*cols + f"{Colors.RESET}")


def main():
    """Main function"""
    db_path = "project_outlines.db"
    
    clear_screen()
    
    if not os.path.exists(db_path):
        print_header("EXPORT PROJECT")
        print(f"\n{Colors.RED}Error:{Colors.RESET} Database file '{db_path}' not found.")
        print(f"{Colors.DIM}Make sure you're running this script in the same directory as your database.{Colors.RESET}\n")
        sys.exit(1)
    
    projects = list_projects(db_path)
    
    if not projects:
        print_header("EXPORT PROJECT")
        print(f"\n{Colors.DIM}(No projects found in the database){Colors.RESET}\n")
        sys.exit(1)
    
    print_header("EXPORT PROJECT")
    
    print(f"\n{Colors.BRIGHT_CYAN}Available Projects:{Colors.RESET}\n")
    for idx, (proj_id, proj_name, created_at) in enumerate(projects, 1):
        print(f"  {Colors.BRIGHT_YELLOW}{idx}{Colors.RESET}. {Colors.BRIGHT_WHITE}{proj_name}{Colors.RESET} {Colors.DIM}(Created: {created_at}){Colors.RESET}")
    
    # Select project
    while True:
        try:
            choice = input(f"\n{Colors.BRIGHT_GREEN}> Enter project number:{Colors.RESET} ").strip()
            choice_num = int(choice)
            
            if 1 <= choice_num <= len(projects):
                project_id = projects[choice_num - 1][0]
                project_name = projects[choice_num - 1][1]
                break
            else:
                print(f"{Colors.RED}Invalid project number. Please try again.{Colors.RESET}")
        except ValueError:
            print(f"{Colors.RED}Invalid input. Please enter a number.{Colors.RESET}")
    
    # Ask for export format
    print(f"\n{Colors.BRIGHT_CYAN}Export Format:{Colors.RESET}\n")
    print(f"  {Colors.BRIGHT_YELLOW}1{Colors.RESET}. Text file only (.txt)")
    print(f"  {Colors.BRIGHT_YELLOW}2{Colors.RESET}. Word document only (.docx)")
    print(f"  {Colors.BRIGHT_YELLOW}3{Colors.RESET}. Both text and Word")
    
    while True:
        format_choice = input(f"\n{Colors.BRIGHT_GREEN}> Select format:{Colors.RESET} ").strip()
        if format_choice in ['1', '2', '3']:
            break
        print(f"{Colors.RED}Invalid choice. Please enter 1, 2, or 3.{Colors.RESET}")
    
    # Get filename
    default_filename = project_name.replace(' ', '_')
    output_filename = input(f"\n{Colors.BRIGHT_GREEN}> Enter filename (default: {Colors.BRIGHT_WHITE}{default_filename}{Colors.BRIGHT_GREEN}):{Colors.RESET} ").strip()
    
    if not output_filename:
        output_filename = default_filename
    
    # Create export folder
    export_dir = ensure_export_folder()
    
    # Get project content
    project_name, content = get_project_content(db_path, project_id)
    
    if not project_name:
        print(f"\n{Colors.RED}Error:{Colors.RESET} Project not found.")
        sys.exit(1)
    
    success_count = 0
    
    # Export text file
    if format_choice in ['1', '3']:
        txt_file = os.path.join(export_dir, f"{output_filename}.txt")
        print(f"\n{Colors.DIM}Exporting to text file...{Colors.RESET}")
        if export_to_text(project_name, content, txt_file):
            print(f"{Colors.GREEN}✓{Colors.RESET} Text file saved: {Colors.BRIGHT_WHITE}{txt_file}{Colors.RESET}")
            success_count += 1
    
    # Export Word document
    if format_choice in ['2', '3']:
        docx_file = os.path.join(export_dir, f"{output_filename}.docx")
        print(f"\n{Colors.DIM}Exporting to Word document...{Colors.RESET}")
        try:
            if export_to_docx(project_name, content, docx_file):
                print(f"{Colors.GREEN}✓{Colors.RESET} Word document saved: {Colors.BRIGHT_WHITE}{docx_file}{Colors.RESET}")
                success_count += 1
        except Exception as e:
            print(f"{Colors.RED}✗{Colors.RESET} Word export failed: {e}")
    
    if success_count > 0:
        print(f"\n{Colors.GREEN}✓{Colors.RESET} Export completed successfully!\n")
    else:
        print(f"\n{Colors.RED}✗ Export failed.{Colors.RESET}\n")
        sys.exit(1)


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print(f"\n\n{Colors.BRIGHT_CYAN}Export cancelled.{Colors.RESET}\n")
        sys.exit(0)
    except Exception as e:
        print(f"\n{Colors.RED}Error:{Colors.RESET} {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
