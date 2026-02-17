#!/usr/bin/env python3
"""
Export Utilities - Export functionality for text and Word documents
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import get_config


class ExportManager:
    """Manage export operations"""
    
    def __init__(self, project_name=None):
        self.config = get_config()
        self.project_name = project_name
        self.datecode = datetime.now().strftime('%Y%m%d')
    
    def get_export_path(self, filename):
        """Get full export path for a filename in project/datecode folder"""
        if self.project_name:
            # Use project-based folder structure
            project_dir = self.config.get_project_export_path(self.project_name, self.datecode)
            return os.path.join(project_dir, filename)
        else:
            # Fallback to base export directory
            export_dir = self.config.get_export_directory()
            return os.path.join(export_dir, filename)
    
    def export_to_text(self, db, project_id, project_name):
        """
        Export project to plain text file
        Returns: filepath or None on error
        """
        # Set project name for export path
        self.project_name = project_name
        
        # Get structured content
        content = self._get_structured_content(db, project_id)
        if not content:
            return None
        
        # Generate filename
        safe_name = "".join(c for c in project_name if c.isalnum() or c in (' ', '-', '_')).strip()
        filename = f"{safe_name}.txt"
        filepath = self.get_export_path(filename)
        
        # Write to file
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"{project_name}\n")
            f.write("=" * len(project_name) + "\n\n")
            
            for mc_id in sorted(content.keys(), key=lambda x: content[x]['order']):
                mc = content[mc_id]
                f.write(f"{mc['name']}\n")
                f.write("-" * len(mc['name']) + "\n\n")
                
                for sc_id in sorted(mc['subcategories'].keys(), key=lambda x: mc['subcategories'][x]['order']):
                    sc = mc['subcategories'][sc_id]
                    
                    # Only print subheading if it has a name
                    if sc['name']:
                        f.write(f"  {sc['name']}\n\n")
                    
                    # Print sentences
                    for sentence in sc['sentences']:
                        f.write(f"    {sentence}\n\n")
                
                f.write("\n")
        
        return filepath
    
    def export_to_word(self, db, project_id, project_name):
        """
        Export project to APA7-formatted Word document
        Returns: filepath or None on error
        """
        # Set project name for export path
        self.project_name = project_name
        
        # Get structured content
        content = self._get_structured_content(db, project_id)
        if not content:
            return None
        
        # Create document
        doc = Document()
        
        # Set default font to Times New Roman 12pt
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Set paragraph spacing for double-spacing
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 2.0
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        
        # Add content
        for mc_id in sorted(content.keys(), key=lambda x: content[x]['order']):
            mc = content[mc_id]
            
            # Add major category (heading) - plain, left-justified
            p = doc.add_paragraph(mc['name'])
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            for sc_id in sorted(mc['subcategories'].keys(), key=lambda x: mc['subcategories'][x]['order']):
                sc = mc['subcategories'][sc_id]
                
                # Add subcategory (subheading) if it has a name - plain, left-justified
                if sc['name']:
                    p = doc.add_paragraph(sc['name'])
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Add sentences
                for sentence in sc['sentences']:
                    p = doc.add_paragraph(sentence)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Generate filename
        safe_name = "".join(c for c in project_name if c.isalnum() or c in (' ', '-', '_')).strip()
        filename = f"{safe_name}.docx"
        filepath = self.get_export_path(filename)
        
        # Save document
        doc.save(filepath)
        
        return filepath
    
    def _get_structured_content(self, db, project_id):
        """
        Get project content structured for export
        Returns: {mc_id: {'name': str, 'order': int, 'subcategories': {...}}}
        """
        # Get all lines
        lines = db.get_all_lines(project_id)
        
        if not lines:
            return None
        
        # Organize content
        content = {}
        for sentence_id, mc_id, mc_name, sc_id, sc_name, s_content, mc_order, sc_order, s_order in lines:
            if mc_id not in content:
                content[mc_id] = {
                    'name': mc_name,
                    'order': mc_order,
                    'subcategories': {}
                }
            
            if sc_id not in content[mc_id]['subcategories']:
                content[mc_id]['subcategories'][sc_id] = {
                    'name': sc_name,
                    'order': sc_order,
                    'sentences': []
                }
            
            if s_content:
                content[mc_id]['subcategories'][sc_id]['sentences'].append(s_content)
        
        return content
